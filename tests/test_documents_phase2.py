from io import BytesIO
from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient


@pytest.fixture()
def client(tmp_path, monkeypatch):
    from app.core.config import get_settings
    from app.database.session import configure_database
    from app.main import create_app

    database_path = tmp_path / "phase2.db"
    upload_dir = tmp_path / "uploads"
    monkeypatch.setenv("UPLOAD_DIR", str(upload_dir))
    monkeypatch.setenv("MAX_UPLOAD_SIZE_BYTES", "1048576")
    get_settings.cache_clear()
    configure_database(f"sqlite:///{database_path}")

    with TestClient(create_app()) as test_client:
        yield test_client

    get_settings.cache_clear()


def _auth_headers(client: TestClient) -> dict[str, str]:
    client.post(
        "/register",
        json={
            "name": "Ada Lovelace",
            "email": "ada@example.com",
            "password": "StrongPass123",
        },
    )
    response = client.post(
        "/login",
        json={"email": "ada@example.com", "password": "StrongPass123"},
    )
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_heading("Company Policy", level=1)
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    return document.tobytes()


def _blank_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page()
    return document.tobytes()


def _watermark_only_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    for y in range(72, 360, 24):
        page.insert_text((72, y), "Scanned by CamScanner")
    return document.tobytes()


def test_upload_docx_lists_metadata_details_and_deletes_file(client, tmp_path):
    headers = _auth_headers(client)
    content = _docx_bytes("Leave policy allows 18 paid days each year.")

    upload_response = client.post(
        "/upload",
        headers=headers,
        files={
            "file": (
                "leave-policy.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload_response.status_code == 201
    uploaded = upload_response.json()
    assert uploaded["filename"] == "leave-policy.docx"
    assert uploaded["file_type"] == "docx"
    assert uploaded["category"] == "Policy"
    assert uploaded["number_of_pages"] == 1

    stored_path = Path(uploaded["filepath"])
    assert stored_path.exists()
    assert tmp_path in stored_path.parents

    list_response = client.get("/documents", headers=headers)
    assert list_response.status_code == 200
    assert [item["filename"] for item in list_response.json()] == ["leave-policy.docx"]

    detail_response = client.get(f"/document/{uploaded['id']}", headers=headers)
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert "Leave policy allows 18 paid days each year." in detail["extracted_text"]

    metadata_response = client.get(f"/metadata/{uploaded['id']}", headers=headers)
    assert metadata_response.status_code == 200
    metadata = metadata_response.json()
    assert metadata["filename"] == "leave-policy.docx"
    assert metadata["file_type"] == "docx"
    assert metadata["number_of_pages"] == 1
    assert metadata["category"] == "Policy"
    assert "upload_date" in metadata

    delete_response = client.delete(f"/document/{uploaded['id']}", headers=headers)
    assert delete_response.status_code == 204
    assert not stored_path.exists()
    assert client.get("/documents", headers=headers).json() == []


def test_upload_pdf_extracts_text_and_page_count(client):
    headers = _auth_headers(client)
    content = _pdf_bytes("Invoice total is 1200 dollars.")

    upload_response = client.post(
        "/upload",
        headers=headers,
        files={"file": ("invoice.pdf", content, "application/pdf")},
    )

    assert upload_response.status_code == 201
    uploaded = upload_response.json()
    assert uploaded["file_type"] == "pdf"
    assert uploaded["number_of_pages"] == 1

    detail_response = client.get(f"/document/{uploaded['id']}", headers=headers)
    assert detail_response.status_code == 200
    assert "Invoice total is 1200 dollars." in detail_response.json()["extracted_text"]


def test_upload_rejects_pdf_without_readable_text(client, tmp_path):
    headers = _auth_headers(client)

    response = client.post(
        "/upload",
        headers=headers,
        files={"file": ("scan.pdf", _blank_pdf_bytes(), "application/pdf")},
    )

    assert response.status_code == 400
    assert "No readable text could be extracted" in response.json()["detail"]
    assert client.get("/documents", headers=headers).json() == []
    assert list((tmp_path / "uploads").glob("*.pdf")) == []


def test_upload_rejects_pdf_with_only_scanner_watermark(client):
    headers = _auth_headers(client)

    response = client.post(
        "/upload",
        headers=headers,
        files={
            "file": (
                "camscanner-watermark.pdf",
                _watermark_only_pdf_bytes(),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 400
    assert "No readable text could be extracted" in response.json()["detail"]
    assert client.get("/documents", headers=headers).json() == []


def test_upload_rejects_missing_auth_unsupported_type_and_large_file(client):
    headers = _auth_headers(client)

    unauthenticated_response = client.post(
        "/upload",
        files={"file": ("notes.txt", b"plain text", "text/plain")},
    )
    unsupported_response = client.post(
        "/upload",
        headers=headers,
        files={"file": ("notes.txt", b"plain text", "text/plain")},
    )
    large_response = client.post(
        "/upload",
        headers=headers,
        files={"file": ("large.pdf", b"x" * (1048576 + 1), "application/pdf")},
    )

    assert unauthenticated_response.status_code in {401, 403}
    assert unsupported_response.status_code == 400
    assert unsupported_response.json()["detail"] == "Only PDF and DOCX files are supported"
    assert large_response.status_code == 413
    assert large_response.json()["detail"] == "File exceeds the maximum allowed size"
