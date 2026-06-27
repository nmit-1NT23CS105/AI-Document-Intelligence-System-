from io import BytesIO

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient


@pytest.fixture()
def client(tmp_path, monkeypatch):
    from app.core.config import get_settings
    from app.database.session import configure_database
    from app.main import create_app

    database_path = tmp_path / "phase4.db"
    upload_dir = tmp_path / "uploads"
    monkeypatch.setenv("UPLOAD_DIR", str(upload_dir))
    monkeypatch.setenv("MAX_UPLOAD_SIZE_BYTES", "1048576")
    monkeypatch.setenv("EMBEDDING_PROVIDER", "local")
    monkeypatch.setenv("VECTOR_STORE_PROVIDER", "database")
    get_settings.cache_clear()
    configure_database(f"sqlite:///{database_path}")

    with TestClient(create_app()) as test_client:
        yield test_client

    get_settings.cache_clear()


def _auth_headers(client: TestClient, email: str = "ada@example.com") -> dict[str, str]:
    client.post(
        "/register",
        json={
            "name": email.split("@")[0].title(),
            "email": email,
            "password": "StrongPass123",
        },
    )
    response = client.post(
        "/login",
        json={"email": email, "password": "StrongPass123"},
    )
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _upload_docx(client: TestClient, headers: dict[str, str], filename: str, text: str):
    return client.post(
        "/upload",
        headers=headers,
        files={
            "file": (
                filename,
                _docx_bytes(text),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def _upload_docx_with_heading(
    client: TestClient,
    headers: dict[str, str],
    filename: str,
    heading: str,
    text: str,
):
    document = DocxDocument()
    document.add_heading(heading, level=1)
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return client.post(
        "/upload",
        headers=headers,
        files={
            "file": (
                filename,
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def test_chat_answers_from_retrieved_document_content_and_saves_history(client):
    headers = _auth_headers(client)
    upload_response = _upload_docx(
        client,
        headers,
        "leave-policy.docx",
        "Employees receive 18 paid leave days every calendar year. "
        "Unused leave expires on March 31.",
    )

    chat_response = client.post(
        "/chat",
        headers=headers,
        json={
            "question": "How many paid leave days do employees receive?",
            "document_id": upload_response.json()["id"],
        },
    )

    assert chat_response.status_code == 200
    body = chat_response.json()
    assert body["answer"] == "Employees receive 18 paid leave days every calendar year."
    assert body["citations"] == [
        {
            "document_id": upload_response.json()["id"],
            "filename": "leave-policy.docx",
            "chunk_index": 0,
        }
    ]

    from app.database.models import ChatHistory
    from app.database.session import SessionLocal

    with SessionLocal() as session:
        history = session.query(ChatHistory).one()

    assert history.question == "How many paid leave days do employees receive?"
    assert history.answer == body["answer"]
    assert history.document_id == upload_response.json()["id"]


def test_chat_refuses_to_answer_when_retrieved_content_does_not_contain_answer(client):
    headers = _auth_headers(client)
    _upload_docx(
        client,
        headers,
        "invoice.docx",
        "Invoice total is 1200 dollars and payment is due next month.",
    )

    response = client.post(
        "/chat",
        headers=headers,
        json={"question": "How many paid leave days are available?", "limit": 3},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["answer"] == "I could not find an answer in the uploaded documents."
    assert body["citations"] == []


def test_chat_document_scope_cannot_access_another_users_document(client):
    ada_headers = _auth_headers(client, "ada@example.com")
    bob_headers = _auth_headers(client, "bob@example.com")
    uploaded = _upload_docx(
        client,
        ada_headers,
        "contract.docx",
        "The contract renewal date is June 30.",
    )

    response = client.post(
        "/chat",
        headers=bob_headers,
        json={
            "question": "What is the renewal date?",
            "document_id": uploaded.json()["id"],
        },
    )

    assert response.status_code == 404
    assert response.json()["detail"] == "Document not found"


def test_chat_answer_ignores_unpunctuated_document_heading(client):
    headers = _auth_headers(client)
    upload_response = _upload_docx_with_heading(
        client,
        headers,
        "runtime-policy.docx",
        "Runtime QA Policy",
        "Employees receive 18 paid leave days every calendar year. "
        "Unused leave expires on March 31.",
    )

    chat_response = client.post(
        "/chat",
        headers=headers,
        json={
            "question": "How many paid leave days do employees receive?",
            "document_id": upload_response.json()["id"],
        },
    )

    assert chat_response.status_code == 200
    assert chat_response.json()["answer"] == (
        "Employees receive 18 paid leave days every calendar year."
    )
