from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument

NO_READABLE_TEXT_MESSAGE = (
    "No readable text could be extracted from this document. "
    "If it is a scanned PDF, install Tesseract OCR or upload a text-based PDF/DOCX."
)

SCANNER_ARTIFACT_PATTERNS = [
    re.compile(r"\bscanned\s+by\s+camscanner\b", re.IGNORECASE),
    re.compile(r"\bcamscanner\b", re.IGNORECASE),
]

TOKEN_PATTERN = re.compile(r"[A-Za-z0-9]+")


class TextExtractionError(ValueError):
    pass


def clean_text(text: str) -> str:
    without_nulls = text.replace("\x00", " ")
    lines = [
        re.sub(r"[ \t\r\f\v]+", " ", line).strip()
        for line in without_nulls.splitlines()
    ]
    normalized = "\n".join(line for line in lines if line)
    if normalized:
        return normalized
    return re.sub(r"\s+", " ", without_nulls).strip()


def _remove_scanner_artifacts(text: str) -> str:
    cleaned = text
    for pattern in SCANNER_ARTIFACT_PATTERNS:
        cleaned = pattern.sub(" ", cleaned)
    return clean_text(cleaned)


def _has_meaningful_text(text: str) -> bool:
    text_without_artifacts = _remove_scanner_artifacts(text)
    tokens = TOKEN_PATTERN.findall(text_without_artifacts)
    return len(tokens) >= 3 and len(set(token.lower() for token in tokens)) >= 3


def _ocr_page_text(page: fitz.Page) -> str:
    try:
        text_page = page.get_textpage_ocr()
        return page.get_text("text", textpage=text_page)
    except Exception:
        return ""


def extract_text_from_pdf(file_path: Path) -> tuple[str, int]:
    with fitz.open(file_path) as document:
        page_text = [page.get_text("text") for page in document]
        extracted_text = clean_text("\n".join(page_text))

        if not _has_meaningful_text(extracted_text):
            ocr_text = clean_text("\n".join(_ocr_page_text(page) for page in document))
            if _has_meaningful_text(ocr_text):
                extracted_text = ocr_text

        extracted_text = _remove_scanner_artifacts(extracted_text)
        if not _has_meaningful_text(extracted_text):
            raise TextExtractionError(NO_READABLE_TEXT_MESSAGE)

        return extracted_text, document.page_count


def extract_text_from_docx(file_path: Path) -> tuple[str, int]:
    document = DocxDocument(file_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    page_count = max(1, len(document.sections))
    extracted_text = clean_text("\n".join(parts))
    if not _has_meaningful_text(extracted_text):
        raise TextExtractionError(NO_READABLE_TEXT_MESSAGE)
    return extracted_text, page_count


def extract_text(file_path: Path, file_type: str) -> tuple[str, int]:
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    if file_type == "docx":
        return extract_text_from_docx(file_path)
    raise ValueError("Unsupported document type")
